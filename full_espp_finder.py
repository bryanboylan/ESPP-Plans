##import python libraries
import docx
from urllib.request import urlopen
import urllib.request
from bs4 import BeautifulSoup
import html

##open new document and set to doc variable
doc = docx.Document()

cik = []
reports = []
tenK = []
tenKurl = []


sp500wiki = "https://en.wikipedia.org/wiki/List_of_S%26P_500_companies"
spList = BeautifulSoup(urllib.request.urlopen(sp500wiki), 'html.parser')
table = spList.find('table').findAll('tr')
table.pop(0)


for row in table:
    symbol = row.find('td').a.string

    for cols in table:
        cols = row.find_all('td')

        cols=[x.text.strip() for x in cols if '000' in x.text.strip()]

    for link in row.find_all('a'):

        link = link.get('href')

        if 'https://www.sec.gov/' in link:
            print(symbol)
            print (link)
            print(cols[0])
            thisCik = cols[0]

            my10kPage = 'https://www.sec.gov/cgi-bin/browse-edgar?type=10-k&CIK=' + thisCik

            ##my10kPage = 'https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK='
            ##+ thisCik + '&type=10-k'

            print(my10kPage)



            try:
                url = my10kPage
                soup = BeautifulSoup(urllib.request.urlopen(url), 'html.parser')
                table = soup.find("table", {"class": "tableFile2"}).find_all('tr')
                href = table[1].a['href']
                url = 'https://www.sec.gov' + href
                soup = BeautifulSoup(urllib.request.urlopen(url), 'html.parser')
                table = soup.find('table').find_all('tr')
                url = 'https://www.sec.gov' + table[1].a['href']
                ##tenKurl.append(url)
                print (url)

               

            except Exception as e:
                break

            page = urlopen(url)

        

            ##create a Beuatiful soup object with page variable. Set to soup variable
            soup = BeautifulSoup(page, 'html.parser')

            ##set espp to the text
            espp = "Employee Stock Purchase Plan" or "espp" or "Employee Stock Option Plan"
            discount = "% discount"
            lower = "lower of"
            stock = "common stock"

            ##declare paragraphs variable and set it to all div text
            paragraphs = soup.find_all('div')

            ##loop through paragraphs/divs
            for ptext in paragraphs:
                ##if espp variable is found in paragraph/div
                if espp in ptext.text and lower in ptext.text:
                    ##print the text of that paragraph/div
                    print (ptext.text)
                    ##easily show the end of a paragraph
                    print("END OF PARAGRAPH\n")
                    ##easily show the end of a paragraph
                    doc.add_paragraph(symbol)
                    doc.add_paragraph(ptext.text)
                    doc.add_paragraph("END OF PARAGRAPH\n")
                elif espp in ptext.text and discount in ptext.text:
                    ##print the text of that paragraph/div
                    print (ptext.text)
                    ##easily show the end of a paragraph
                    print("END OF PARAGRAPH\n")
                    ##add paragraph to word doc
                    doc.add_paragraph(symbol)
                    doc.add_paragraph(ptext.text)
                    doc.add_paragraph("END OF PARAGRAPH\n")
                elif espp in ptext.text and stock in ptext.text:
                    print (ptext.text)
                    ##easily show the end of a paragraph
                    print("END OF PARAGRAPH\n")
                    ##add paragraph to word doc
                    doc.add_paragraph(symbol)
                    doc.add_paragraph(ptext.text)
                    doc.add_paragraph("END OF PARAGRAPH\n")


            

##save word doc
doc.save('10k espp.docx')

